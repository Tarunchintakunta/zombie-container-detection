"""
Generate IEEE-formatted .docx documents for Paper 1 and Paper 2 critical reviews.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))


def set_narrow_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(1.91)
        section.right_margin = Cm(1.91)


def setup_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    pf.first_line_indent = Cm(0.75)

    # Title style
    if 'IEEE Title' not in [s.name for s in doc.styles]:
        title_style = doc.styles.add_style('IEEE Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 0, 0)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_before = Pt(0)
        title_style.paragraph_format.space_after = Pt(12)
        title_style.paragraph_format.first_line_indent = Cm(0)

    # Author style
    if 'IEEE Author' not in [s.name for s in doc.styles]:
        author_style = doc.styles.add_style('IEEE Author', WD_STYLE_TYPE.PARAGRAPH)
        author_style.font.name = 'Times New Roman'
        author_style.font.size = Pt(11)
        author_style.font.color.rgb = RGBColor(0, 0, 0)
        author_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_style.paragraph_format.space_before = Pt(0)
        author_style.paragraph_format.space_after = Pt(4)
        author_style.paragraph_format.first_line_indent = Cm(0)

    # Section heading style
    if 'IEEE Heading' not in [s.name for s in doc.styles]:
        heading_style = doc.styles.add_style('IEEE Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Times New Roman'
        heading_style.font.size = Pt(10)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)
        heading_style.paragraph_format.first_line_indent = Cm(0)
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Reference style
    if 'IEEE Reference' not in [s.name for s in doc.styles]:
        ref_style = doc.styles.add_style('IEEE Reference', WD_STYLE_TYPE.PARAGRAPH)
        ref_style.font.name = 'Times New Roman'
        ref_style.font.size = Pt(8)
        ref_style.font.color.rgb = RGBColor(0, 0, 0)
        ref_style.paragraph_format.space_before = Pt(0)
        ref_style.paragraph_format.space_after = Pt(3)
        ref_style.paragraph_format.first_line_indent = Cm(0)
        ref_style.paragraph_format.left_indent = Cm(0.5)


def add_first_paragraph(doc, text):
    """Add a paragraph with no first-line indent (IEEE first paragraph of section)."""
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def generate_paper1():
    doc = Document()
    set_narrow_margins(doc)
    setup_styles(doc)

    # Title
    doc.add_paragraph(
        'Critical Review: Enhancing Kubernetes Resilience through Anomaly Detection and Prediction',
        style='IEEE Title'
    )

    # Author block
    doc.add_paragraph(
        'Prepared for: Heuristic-Based Approach to Detect Zombie Containers in Kubernetes for Resource Optimization',
        style='IEEE Author'
    )
    doc.add_paragraph(
        'Anurag Baiju — 23409223 — MSc in Cloud Computing — National College of Ireland',
        style='IEEE Author'
    )

    # Paper details heading
    doc.add_paragraph('I. PAPER UNDER REVIEW', style='IEEE Heading')

    add_first_paragraph(doc,
        'V. Anemogiannis, C. Andreou, K. Myrtollari, K. Panagidi, and S. Hadjiefthymiades, '
        '"Enhancing Kubernetes Resilience through Anomaly Detection and Prediction," '
        'arXiv preprint arXiv:2503.14114, March 2025. National and Kapodistrian University of Athens, Greece. '
        'EU Horizon Europe EO4EU project. Available: https://arxiv.org/abs/2503.14114'
    )

    # Section II
    doc.add_paragraph('II. SUMMARY OF THE PAPER', style='IEEE Heading')

    add_first_paragraph(doc,
        'Anemogiannis et al. (2025) present a framework for enhancing the resilience of Kubernetes clusters '
        'through anomaly detection and prediction, developed as part of the European Union Horizon Europe EO4EU '
        'Earth-observation project at the National and Kapodistrian University of Athens. The framework encodes '
        'the full Kubernetes resource hierarchy \u2014 Nodes, Deployments, ReplicaSets, Namespaces, and Pods \u2014 '
        'as a heterogeneous directed graph stored in a Neo4j graph database, where edges capture ownership and '
        'co-location relationships and node attributes are populated from Prometheus telemetry. Prometheus, '
        'configured with a fifteen-second scrape interval, collects CPU saturation, memory utilisation, and '
        'service request-rate metrics from every instrumented container in the cluster; Thanos extends Prometheus '
        'with long-term storage and high availability, and Grafana provides supplementary visualisation. The '
        'detection pipeline operates in two sequential phases. In the unsupervised phase, three algorithms \u2014 '
        'Isolation Forest, Density-Based Spatial Clustering of Applications with Noise (DBSCAN), and One-Class '
        'Support Vector Machine (OCSVM) \u2014 are evaluated to establish a dynamic boundary of normality against '
        'which future observations can be compared. In the supervised phase, three classifiers \u2014 Decision Trees, '
        'Support Vector Machines, and Logistic Regression \u2014 are trained on the labelled output of the '
        'unsupervised phase to predict whether incoming metric samples represent anomalous cluster behaviour. '
        'Hyperparameter optimisation is performed using Optuna, and the entire pipeline is evaluated over '
        'N equal to one hundred experimental iterations on the production EO4EU cluster, with two controlled '
        'stress-test scenarios: CPU stress induced by thirty-two concurrent workers executing square-root '
        'computations, and memory stress induced by thirty-two concurrent workers executing malloc allocations. '
        'The best unsupervised result is a silhouette score of 0.794 achieved by Isolation Forest, while the '
        'best supervised result is an F1 score of 0.886 achieved by Decision Trees, with Support Vector Machines '
        'reaching 0.878 and Logistic Regression reaching 0.870.'
    )

    # Section III
    doc.add_paragraph('III. RELEVANCE TO THE PRESENT RESEARCH', style='IEEE Heading')

    add_first_paragraph(doc,
        'The relevance of this work to the present thesis is substantial and operates on three levels. First, '
        'the monitoring infrastructure that Anemogiannis et al. employ \u2014 Prometheus collecting CPU and memory '
        'metrics at fifteen-second intervals from a Kubernetes cluster \u2014 is architecturally identical to the '
        'monitoring stack adopted in this thesis, where Prometheus scrapes container_cpu_usage_seconds_total, '
        'container_memory_usage_bytes, container_network_receive_bytes_total, and '
        'container_network_transmit_bytes_total at the same fifteen-second resolution within a Minikube cluster. '
        'This congruence is methodologically significant because it permits a controlled comparison between the '
        'machine-learning anomaly scores reported by Anemogiannis et al. and the heuristic composite scores '
        'produced by the five-rule detector developed in this thesis, without introducing confounds arising from '
        'differences in metric granularity or collection infrastructure. Second, the Isolation Forest results '
        'reported in that paper \u2014 silhouette score of 0.794 \u2014 provide a concrete, published, quantitative '
        'baseline against which the heuristic approach can be benchmarked. This is particularly important because '
        'the evaluation plan for this thesis includes a comparative analysis against machine-learning techniques, '
        'with Isolation Forest identified as the primary comparator following the precedent established by Cui et '
        'al. (2024) for Docker container anomaly detection. Having an independent, recent set of Isolation Forest '
        'results obtained on a Kubernetes cluster using the same monitoring stack strengthens the validity of any '
        'performance comparison. Third, the operational overhead disclosed in the Anemogiannis et al. framework '
        '\u2014 the requirement for a Neo4j graph database to store cluster topology, periodic graph-update cycles '
        'to refresh node attributes, Optuna-based hyperparameter sweeps, and model-retraining procedures to '
        'accommodate concept drift \u2014 provides a concrete point of contrast with the heuristic detector developed '
        'here, which requires no external database, no training data, no hyperparameter optimisation, and no '
        'retraining cycle, operating instead through deterministic rule evaluation over live Prometheus queries.'
    )

    # Section IV
    doc.add_paragraph('IV. IDENTIFIED GAPS AND CRITICAL ANALYSIS', style='IEEE Heading')

    add_first_paragraph(doc,
        'Despite these strengths, a critical examination of Anemogiannis et al. (2025) reveals five specific '
        'gaps that limit its applicability to the zombie container detection problem and that the present thesis '
        'is designed to address. The first and most fundamental gap is that the framework is engineered to detect '
        'active performance degradation \u2014 that is, anomalous upward deviations in CPU or memory consumption '
        'relative to learned baselines \u2014 and not the inverse pattern that characterises zombie containers. A '
        'zombie container, as defined in this thesis, exhibits sustained near-zero CPU utilisation (below five '
        'percent for thirty or more minutes), stable or slowly increasing memory allocation, and negligible '
        'network activity; this behavioural profile sits comfortably within the statistical norm of a '
        'heterogeneous cluster and would therefore be classified as healthy, not anomalous, by any of the six '
        'machine-learning models evaluated by Anemogiannis et al. The Isolation Forest algorithm, which scores '
        'each observation by the number of random splits required to isolate it, assigns high anomaly scores to '
        'points that are easily separated from the majority; a container with near-zero CPU and low memory churn '
        'is, by definition, close to the centre of the normal distribution and would receive a low anomaly score. '
        'The heuristic detector developed in this thesis addresses this inversion directly: Rule 1 (sustained '
        'low CPU with memory allocation, weighted at thirty-five percent) triggers precisely when CPU remains '
        'below five percent for more than thirty minutes while memory remains stable or increases, a combination '
        'that machine-learning normality models would consider unremarkable.'
    )

    doc.add_paragraph(
        'The second gap is the absence of any discussion of zombie containers, idle containers, orphan resources, '
        'or resource waste in the paper. A full-text examination of the twenty-seven-page manuscript confirms '
        'that the terms "zombie," "idle," "orphan," "stale," and "resource waste" do not appear. The problem '
        'formulation centres on cluster resilience and the propagation of anomalies across the Kubernetes '
        'resource graph, not on the identification of containers that consume resources without performing '
        'productive work. This omission is not a criticism of the paper\u2019s quality \u2014 it simply addresses a '
        'different problem \u2014 but it establishes that zombie container detection remains an unoccupied research '
        'niche even after the publication of this recent and methodologically rigorous work.'
    )

    doc.add_paragraph(
        'The third gap concerns interpretability and operational governance. The Decision Tree classifier that '
        'achieves F1 of 0.886 is, at sufficient depth, not readily interpretable by a site-reliability engineer '
        'who needs to understand why a particular container was flagged and whether the flag warrants remedial '
        'action. In production Kubernetes environments, where false positives can trigger unnecessary pod '
        'evictions and service disruptions, the ability to trace a detection decision back to a specific, '
        'human-readable rule \u2014 for example, "CPU below five percent for forty-two minutes with memory stable '
        'at one hundred and twenty megabytes and no network traffic" \u2014 is a governance requirement that opaque '
        'machine-learning models cannot satisfy without additional explainability tooling. Each of the five '
        'heuristic rules implemented in this thesis produces a per-rule score between zero and one together '
        'with a structured details dictionary that records the exact metric values and durations that '
        'contributed to the score, enabling engineers to audit every detection decision at the level of '
        'individual threshold comparisons.'
    )

    doc.add_paragraph(
        'The fourth gap is the infrastructure overhead introduced by the graph-based approach. Maintaining a '
        'Neo4j database, running periodic graph-update jobs, and executing Optuna hyperparameter searches '
        'represent ongoing operational costs in terms of compute, storage, and engineering attention. For '
        'small-to-medium clusters \u2014 and in particular for the Minikube-based development and staging '
        'environments that are the experimental setting of this thesis \u2014 such overhead may exceed the overhead '
        'of the very problem being detected. The heuristic detector, by contrast, runs as a single Python '
        'process within a Kubernetes Deployment, requesting one hundred millicores of CPU and two hundred and '
        'fifty-six megabytes of memory, and its only external dependency is a reachable Prometheus endpoint.'
    )

    doc.add_paragraph(
        'The fifth gap is the absence of a comparative evaluation against simpler detection methods. '
        'Anemogiannis et al. compare six machine-learning algorithms against one another \u2014 Isolation Forest '
        'versus DBSCAN versus OCSVM versus Decision Trees versus SVM versus Logistic Regression \u2014 but never '
        'benchmark any of them against a static-threshold alert (for example, CPU below five percent for more '
        'than thirty minutes) or a rule-based heuristic system. This omission leaves an important empirical '
        'question unanswered: for the specific sub-problem of identifying containers that have ceased useful '
        'work, is the full weight of a machine-learning pipeline necessary, or can a well-designed set of '
        'heuristic rules achieve comparable or superior performance at a fraction of the cost? The evaluation '
        'plan for this thesis is designed to answer this question by comparing the heuristic detector against '
        'both a simple static-threshold baseline and the Isolation Forest results reported by Anemogiannis et '
        'al., measuring accuracy, precision, recall, F1 score, false positive rate, detection latency, and '
        'computational overhead across controlled test scenarios that include normal containers, legitimately '
        'idle containers (batch jobs waiting for input, cron jobs between scheduled runs), and five distinct '
        'zombie container archetypes (sustained low CPU, memory leak, stuck process, network timeout, and '
        'resource ratio imbalance).'
    )

    # Section V
    doc.add_paragraph('V. HOW THE PRESENT RESEARCH FILLS THE GAPS', style='IEEE Heading')

    add_first_paragraph(doc,
        'In summary, Anemogiannis et al. (2025) contribute a rigorous, well-evaluated framework for '
        'machine-learning-based anomaly detection in Kubernetes using the same Prometheus-based monitoring '
        'infrastructure adopted in this thesis. Their Isolation Forest and Decision Tree results provide '
        'valuable quantitative benchmarks. However, their framework targets active performance anomalies '
        'rather than zombie containers, operates as an opaque machine-learning system with significant '
        'infrastructure overhead, and does not compare against simpler detection methods. The heuristic-based '
        'approach developed in this thesis is positioned to fill these gaps by providing a lightweight, '
        'transparent, purpose-built zombie detection mechanism that can be deployed alongside existing '
        'monitoring infrastructure, evaluated against established machine-learning baselines, and governed '
        'through human-readable rules that align with the operational realities of Kubernetes cluster management. '
        'The five heuristic rules \u2014 sustained low CPU with memory allocation (thirty-five percent weight), '
        'memory leak pattern (twenty-five percent), stuck process pattern (fifteen percent), network timeout '
        'pattern (fifteen percent), and resource ratio imbalance (ten percent) \u2014 produce a composite zombie '
        'score on a zero-to-one-hundred scale, with classification thresholds of seventy for confirmed zombie '
        'and forty for potential zombie. This scoring system is fully configurable, enabling operations teams '
        'to tune the precision\u2013recall trade-off for their specific cluster environment without retraining a '
        'model or modifying algorithmic internals.'
    )

    # References
    doc.add_paragraph('REFERENCES', style='IEEE Heading')

    refs = [
        '[1] V. Anemogiannis, C. Andreou, K. Myrtollari, K. Panagidi, and S. Hadjiefthymiades, '
        '"Enhancing Kubernetes Resilience through Anomaly Detection and Prediction," '
        'arXiv preprint arXiv:2503.14114, Mar. 2025.',

        '[2] G. Cui et al., "A Docker Container Anomaly Monitoring System Based on Optimized '
        'Isolation Forest," IEEE Access, 2024.',

        '[3] H. Li, W. Rao, B. Hu, Y. Tian, and J. Shen, "Energy-aware elastic scaling algorithm '
        'for microservices in Kubernetes clouds," J. Netw. Comput. Appl., Jun. 2025.',
    ]
    for ref in refs:
        doc.add_paragraph(ref, style='IEEE Reference')

    out = os.path.join(DOCS_DIR, 'Paper1_Anemogiannis_Critical_Review.docx')
    doc.save(out)
    print(f"Saved: {out}")


def generate_paper2():
    doc = Document()
    set_narrow_margins(doc)
    setup_styles(doc)

    # Title
    doc.add_paragraph(
        'Critical Review: Energy-Aware Elastic Scaling Algorithm for Microservices in Kubernetes Clouds',
        style='IEEE Title'
    )

    # Author block
    doc.add_paragraph(
        'Prepared for: Heuristic-Based Approach to Detect Zombie Containers in Kubernetes for Resource Optimization',
        style='IEEE Author'
    )
    doc.add_paragraph(
        'Anurag Baiju — 23409223 — MSc in Cloud Computing — National College of Ireland',
        style='IEEE Author'
    )

    # Paper details
    doc.add_paragraph('I. PAPER UNDER REVIEW', style='IEEE Heading')

    add_first_paragraph(doc,
        'H. Li, W. Rao, B. Hu, Y. Tian, and J. Shen, "Energy-aware elastic scaling algorithm for '
        'microservices in Kubernetes clouds," Journal of Network and Computer Applications (Elsevier), '
        'June 2025. Chongqing University of Posts and Telecommunications, China. Peer-reviewed, '
        'impact factor approximately 7.5. Available: '
        'https://www.sciencedirect.com/science/article/abs/pii/S1084804525001158 and '
        'https://www.researchgate.net/publication/392348403'
    )

    # Section II
    doc.add_paragraph('II. SUMMARY OF THE PAPER', style='IEEE Heading')

    add_first_paragraph(doc,
        'Li et al. (2025) address the problem of energy-inefficient elastic scaling in Kubernetes-based '
        'microservice deployments, published in the Journal of Network and Computer Applications (Elsevier), '
        'a peer-reviewed venue with an impact factor of approximately 7.5. The central observation motivating '
        'their work is stated explicitly and bears direct relevance to this thesis: "the default scaling '
        'mechanisms in Kubernetes fail to effectively distinguish and manage resource consumption of idle '
        'containers, leading to resource waste and degraded system performance." The authors further note that '
        '"Kubernetes\u2019 default metrics fail to distinguish between active and idle containers" and that, as a '
        'consequence, "idle containers continue to consume resources, which significantly impacts cold start '
        'latency and overall system performance." These three statements validate the core premise of the '
        'present research \u2014 that containers which consume resources without performing useful work represent a '
        'tangible and under-addressed source of waste in Kubernetes environments \u2014 and establish that even '
        'recent, peer-reviewed research recognises the detection of idle and zombie containers as an open problem.'
    )

    doc.add_paragraph(
        'To address this problem, Li et al. propose an Energy-Aware Elastic Scaling Algorithm (EAES) that '
        'extends the Kubernetes Horizontal Pod Autoscaler with an energy-efficiency model grounded in Service '
        'Level Agreement constraints. The algorithm integrates two control mechanisms. A feedforward controller '
        'anticipates workload demand by modelling expected access arrival rates and container processing '
        'capacities, enabling pre-emptive scaling before demand spikes cause latency violations. A feedback '
        'controller monitors observed utilisation and adjusts the number of active container replicas, '
        'explicitly accounting for the energy cost of idle containers by incorporating CPU power draw during '
        'container runtime, Power Usage Effectiveness (PUE), response-time targets, and container processing '
        'rates into a unified optimisation function. The feedback loop periodically releases excess containers '
        'whose utilisation falls below the threshold implied by the SLA-energy trade-off, thereby reducing the '
        'proportion of cluster resources consumed by idle instances. Experimental evaluation demonstrates a '
        'reduction of approximately 15.34 percent in total energy consumption relative to the default Kubernetes '
        'Horizontal Pod Autoscaler, while maintaining SLA compliance across the tested workload scenarios.'
    )

    # Section III
    doc.add_paragraph('III. RELEVANCE TO THE PRESENT RESEARCH', style='IEEE Heading')

    add_first_paragraph(doc,
        'Three aspects of this work are directly relevant to the present thesis and merit careful analysis. '
        'First, the explicit identification of idle containers as a distinct category of Kubernetes resource '
        'waste validates the problem definition that underpins this research at the level of peer-reviewed, '
        'journal-published evidence. The literature review of this thesis establishes that zombie containers '
        '\u2014 containers that are running but performing no useful work \u2014 represent a significant and '
        'under-detected source of operational cost and carbon footprint in cloud deployments; Li et al. provide '
        'independent, peer-reviewed corroboration of this claim by demonstrating that idle containers cause '
        'measurable energy waste and performance degradation in production-grade Kubernetes clusters. Second, '
        'the authors\u2019 observation that Kubernetes\u2019 built-in metrics cannot differentiate active from idle '
        'containers directly motivates the design of a dedicated detection mechanism. The default Kubernetes '
        'metrics server reports aggregate CPU and memory utilisation per pod but does not expose the temporal '
        'patterns, network activity correlations, or resource-allocation ratios that are necessary to '
        'distinguish a container that is temporarily idle (a batch job waiting for its next input file, a cron '
        'job sleeping between scheduled executions, a cold-standby service maintaining readiness for failover) '
        'from a container that has permanently ceased useful work (a process stuck in an infinite retry loop '
        'against a decommissioned endpoint, a service whose upstream dependency was removed without cascading '
        'deletion, a deployment that was scaled up for a load test and never scaled back down). The heuristic '
        'rules implemented in this thesis are designed precisely to make this distinction by analysing sustained '
        'temporal patterns across multiple metrics rather than relying on instantaneous utilisation snapshots. '
        'Third, the energy-consumption model proposed by Li et al. offers a complementary quantification '
        'framework: once a zombie container is identified and terminated by the heuristic detector, the energy '
        'savings attributable to its removal can, in principle, be estimated using the EAES energy model, '
        'linking the detection contribution of this thesis to the energy-efficiency contribution of Li et al. '
        'in a chain of complementary capabilities.'
    )

    # Section IV
    doc.add_paragraph('IV. IDENTIFIED GAPS AND CRITICAL ANALYSIS', style='IEEE Heading')

    add_first_paragraph(doc,
        'Despite these points of alignment, a critical examination of Li et al. (2025) reveals four specific '
        'gaps that the present thesis is designed to address. The first and most significant gap is that EAES '
        'is fundamentally a scaling algorithm, not a detection or classification mechanism. It adjusts the '
        'number of container replicas in response to the balance between workload demand and energy cost, but '
        'it does not produce a per-container verdict of "zombie," "legitimately idle," or "active." This '
        'distinction is non-trivial and has practical consequences. Consider a Kubernetes namespace containing '
        'three deployments: a web server handling live traffic, a batch processor that runs every six hours and '
        'sleeps between runs, and an orphaned monitoring sidecar whose parent service was deleted three weeks '
        'ago. All three exhibit low CPU utilisation between processing windows, but only the third is a zombie. '
        'The EAES scaling algorithm would treat all three identically \u2014 as candidates for replica reduction '
        'when utilisation falls below the energy-cost threshold \u2014 potentially disrupting the batch processor\u2019s '
        'next scheduled run or the web server\u2019s cold-standby capacity. The heuristic detector developed in '
        'this thesis addresses this problem through five orthogonal rules that examine not just instantaneous '
        'utilisation but sustained behavioural patterns. Rule 1 (sustained low CPU with memory allocation, '
        'weighted at thirty-five percent of the composite score) checks whether CPU has remained below five '
        'percent for at least thirty minutes while memory remains stable or increases and network activity is '
        'negligible; a batch processor that exhibits periodic CPU spikes every six hours would not trigger this '
        'rule because its CPU history contains recent spikes, whereas the orphaned sidecar with flat-line CPU '
        'and no network traffic would score highly. Rule 3 (stuck process pattern, weighted at fifteen percent) '
        'specifically detects the pattern of brief CPU spikes followed by extended idle periods repeated three '
        'or more times, which is characteristic of a process trapped in a retry loop \u2014 a zombie archetype '
        'that a scaling algorithm would not distinguish from normal periodic processing.'
    )

    doc.add_paragraph(
        'The second gap concerns the temporal granularity and analytical depth of the metrics used. EAES '
        'operates on metrics that are processed through a control-loop abstraction \u2014 arrival rates, processing '
        'rates, and response times fed into feedforward and feedback controllers \u2014 and makes scaling decisions '
        'on the basis of short-window aggregates aligned with the control-loop cycle time. It does not perform '
        'the kind of extended temporal pattern analysis that is necessary to identify zombie containers with '
        'high confidence while minimising false positives against legitimately idle workloads. The heuristic '
        'detector developed in this thesis uses a configurable look-back window (defaulting to sixty minutes '
        'but extensible to longer observation periods) with a fifteen-second sampling interval, yielding '
        'approximately two hundred and forty data points per metric per container. This temporal depth enables '
        'Rule 2 (memory leak pattern, weighted at twenty-five percent) to detect gradual memory increases '
        'exceeding five percent over one hour while CPU remains below one percent \u2014 a signature of a zombie '
        'container that is leaking memory through an abandoned data structure or unclosed connection \u2014 and '
        'enables Rule 4 (network timeout pattern, weighted at fifteen percent) to detect periodic, low-volume '
        'network connection attempts with a coefficient of variation below 0.5 in inter-attempt intervals, '
        'characteristic of a container retrying connections to a decommissioned service at regular intervals. '
        'Neither of these patterns is detectable through the aggregate, short-window metrics that feed the '
        'EAES control loop.'
    )

    doc.add_paragraph(
        'The third gap is the metrics source. Li et al. explicitly acknowledge that Kubernetes\u2019 default '
        'metrics are inadequate for distinguishing active from idle containers, yet their EAES algorithm '
        'operates on metrics derived from the Kubernetes metrics pipeline rather than from a dedicated '
        'monitoring system such as Prometheus. The heuristic detector developed in this thesis queries '
        'Prometheus directly using PromQL expressions \u2014 rate(container_cpu_usage_seconds_total[5m]) for CPU '
        'usage rate, container_memory_usage_bytes for memory consumption, '
        'rate(container_network_receive_bytes_total[5m]) and rate(container_network_transmit_bytes_total[5m]) '
        'for network activity, and container_spec_memory_limit_bytes and container_spec_cpu_quota for resource '
        'allocation limits \u2014 obtaining time-series data at the full fifteen-second resolution that Prometheus '
        'provides. This richer telemetry source enables the multi-metric, temporal-pattern analysis that '
        'underpins each of the five heuristic rules, and its adoption is directly motivated by the limitation '
        'that Li et al. themselves identify in Kubernetes\u2019 native metrics.'
    )

    doc.add_paragraph(
        'The fourth gap is that the EAES algorithm is embedded within the scaling control loop and cannot be '
        'deployed as a standalone detection, auditing, or governance tool. In many operational contexts \u2014 '
        'particularly those governed by change-control processes, compliance requirements, or multi-team '
        'ownership structures \u2014 the ability to detect and report zombie containers without automatically '
        'terminating them is a prerequisite for adoption. An operations team may wish to review a list of '
        'suspected zombies, validate the detections against application-specific knowledge (for example, '
        'confirming that a flagged container is not a legitimate cold-standby instance), and then take '
        'deliberate remedial action through established change-management procedures. The heuristic detector '
        'developed in this thesis supports exactly this workflow: it operates as an independent Kubernetes '
        'Deployment that can run in continuous monitoring mode (checking every five minutes by default) or as '
        'a one-time audit, producing structured output in either human-readable text or machine-parseable JSON '
        'format. Each output record includes the container\u2019s namespace, pod name, container name, and node, '
        'together with the composite zombie score on a zero-to-one-hundred scale, the per-rule breakdown '
        'showing which of the five rules triggered and with what individual score, and a details dictionary '
        'for each rule that records the exact metric values and temporal durations that contributed to the '
        'score. This level of transparency enables engineers to trace every detection decision back to '
        'specific, auditable threshold comparisons and to adjust individual rule thresholds or weights to '
        'tune the precision\u2013recall trade-off for their specific operational context.'
    )

    # Section V
    doc.add_paragraph('V. HOW THE PRESENT RESEARCH FILLS THE GAPS', style='IEEE Heading')

    add_first_paragraph(doc,
        'Taken together, the four gaps identified in Li et al. (2025) \u2014 the absence of per-container '
        'classification, the limited temporal analytical depth, the reliance on default Kubernetes metrics '
        'rather than Prometheus, and the tight coupling with the scaling control loop \u2014 define a clear and '
        'defensible space for the contribution of this thesis. Li et al. establish that idle containers are '
        'a real and measurable source of waste in Kubernetes environments, that Kubernetes itself cannot '
        'distinguish active from idle containers, and that addressing this waste yields quantifiable energy '
        'savings. The heuristic detector developed in this thesis builds on these findings by providing the '
        'missing detection and classification layer: a standalone, Prometheus-powered, rule-based mechanism '
        'that examines sustained behavioural patterns across CPU, memory, and network metrics to classify '
        'containers as zombie, potentially zombie, or normal, with full per-rule transparency and no '
        'dependency on the cluster\u2019s scaling infrastructure. Where Li et al. answer the question "how can we '
        'scale more efficiently in the presence of idle containers," this thesis answers the prior question '
        '"which containers are zombies and should be investigated for removal" \u2014 a question that must be '
        'answered before any scaling or termination action can be taken responsibly.'
    )

    doc.add_paragraph(
        'The choice of Li et al. (2025) as an anchoring reference is further justified by the complementary '
        'relationship between their contribution and the contribution of this thesis within the broader '
        'Kubernetes resource optimisation landscape. A complete solution to the idle container problem requires '
        'two capabilities in sequence: first, detecting which containers are truly zombie (the contribution of '
        'this thesis) and, second, acting on that detection through scaling, termination, or remediation (the '
        'domain addressed by Li et al. and by the Kubernetes Horizontal Pod Autoscaler more broadly). By '
        'positioning the heuristic detector as a pre-scaling diagnostic \u2014 a tool that identifies zombie '
        'containers before scaling decisions are made \u2014 this thesis does not compete with the EAES algorithm '
        'but rather complements it, providing the classification intelligence that scaling algorithms assume '
        'but do not themselves generate. This complementary positioning also opens a path for future work in '
        'which the heuristic detector\u2019s output could feed directly into an energy-aware scaling framework such '
        'as EAES, creating an integrated pipeline that first identifies zombies through transparent heuristic '
        'rules and then reclaims their resources through energy-optimised scaling \u2014 a pipeline that would '
        'combine the interpretability and low overhead of heuristics with the energy-modelling sophistication '
        'of control-theoretic scaling.'
    )

    # Section VI
    doc.add_paragraph('VI. LIMITATIONS OF THIS REFERENCE', style='IEEE Heading')

    add_first_paragraph(doc,
        'A limitation of relying on Li et al. (2025) as an anchoring reference is that the full text of the '
        'paper is available only through institutional subscription to the Journal of Network and Computer '
        'Applications on ScienceDirect. This concern is mitigated by two factors: the National College of '
        'Ireland maintains an institutional subscription to Elsevier journals through its library, providing '
        'the author of this thesis with full-text access; and the paper is also listed on ResearchGate, where '
        'the authors may provide access upon request. The peer-reviewed status of the paper, its publication '
        'in a journal with an impact factor of approximately 7.5, and its explicit treatment of idle containers '
        'in Kubernetes environments make it a more authoritative and directly relevant reference for the zombie '
        'container detection problem than the alternative sources available in the 2024 to 2026 academic '
        'literature, none of which address idle or zombie container detection as a primary research problem.'
    )

    # References
    doc.add_paragraph('REFERENCES', style='IEEE Heading')

    refs = [
        '[1] H. Li, W. Rao, B. Hu, Y. Tian, and J. Shen, "Energy-aware elastic scaling algorithm for '
        'microservices in Kubernetes clouds," J. Netw. Comput. Appl., Elsevier, Jun. 2025.',

        '[2] V. Anemogiannis, C. Andreou, K. Myrtollari, K. Panagidi, and S. Hadjiefthymiades, '
        '"Enhancing Kubernetes Resilience through Anomaly Detection and Prediction," '
        'arXiv preprint arXiv:2503.14114, Mar. 2025.',

        '[3] G. Cui et al., "A Docker Container Anomaly Monitoring System Based on Optimized '
        'Isolation Forest," IEEE Access, 2024.',
    ]
    for ref in refs:
        doc.add_paragraph(ref, style='IEEE Reference')

    out = os.path.join(DOCS_DIR, 'Paper2_Li_Energy_Aware_Critical_Review.docx')
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == '__main__':
    generate_paper1()
    generate_paper2()
    print("Done. Both .docx files generated.")
